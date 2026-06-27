"""
clean_tiktok.py
Limpia el extracto de pedidos de TikTok y genera la documentacion para el ERP.

Que hace:
1. Abre una ventana para elegir el archivo (Excel o CSV) descargado de TikTok.
2. Quita la segunda fila de cabecera (la fila de explicaciones).
3. Se queda solo con las columnas necesarias y las renombra en espanol.
4. Limpia los importes (quita "EUR", convierte a numero).
5. Normaliza el pais (Spain/España -> España, Germany/Deutschland -> Alemania, etc.).
6. Calcula Precio_Unitario (Total / Qty / 1.21) e IVA (21% del unitario).
7. Convierte Shipped Time a Fecha_Recogida en formato DD/MM/YYYY.
8. Ordena por Fecha_Recogida, Pais e ID_Pedido.
9. Crea una carpeta "YYYYMMDD_N" junto al original y dentro genera:
     - Orders_invoice.xlsx
     - Resumen_YYYYMMDD.xlsx
     - MANIFIESTO DE ENTREGA_SEUR_YYYYMMDD.docx  (pedidos NO Espana)
     - MANIFIESTO DE ENTREGA_CTT_YYYYMMDD.docx   (pedidos Espana)
"""

import re
import sys
import tkinter as tk
from tkinter import filedialog, messagebox
from pathlib import Path
from datetime import datetime

import pandas as pd

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    Document = None  # se avisa al usuario si falta python-docx


# ----------------------------------------------------------------------
# Configuracion
# ----------------------------------------------------------------------
# nombre en TikTok -> nombre interno
RENOMBRAR = {
    "Order ID": "ID_Pedido",
    "Seller SKU": "Codigo_Producto",
    "Quantity": "Qty",
    "Order Amount": "Total",
    "Shipped Time": "Fecha_Recogida",
    "Country": "Pais",
}

COLUMNAS_NECESARIAS = list(RENOMBRAR.keys())

# columnas de la factura (Orders_invoice), en orden
ORDEN_FINAL = [
    "Pais",
    "ID_Pedido",
    "Fecha_Recogida",
    "Codigo_Producto",
    "Qty",
    "Precio_Unitario",
    "IVA",
    "Total",
]

# normalizacion de pais (clave en minusculas)
MAPA_PAIS = {
    "espana": "España",
    "españa": "España",
    "spain": "España",
    "italy": "Italia",
    "italia": "Italia",
    "germany": "Alemania",
    "deutschland": "Alemania",
    "alemania": "Alemania",
    "ireland": "Irlanda",
    "irlanda": "Irlanda",
}

MESES_ES = [
    "enero", "febrero", "marzo", "abril", "mayo", "junio",
    "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
]

IVA_TIPO = 0.21  # 21%


# ----------------------------------------------------------------------
# Funciones de limpieza
# ----------------------------------------------------------------------
def limpiar_dinero(valor):
    """Convierte 'EUR 12.34', '12,34 EUR', '€12.34' ... en el numero 12.34"""
    if pd.isna(valor):
        return None
    texto = str(valor)
    texto = re.sub(r"(?i)eur", "", texto)
    texto = texto.replace("€", "").replace("\u00a0", "").strip()
    texto = re.sub(r"[^\d,.\-]", "", texto)
    if texto in ("", "-", ".", ","):
        return None
    if "," in texto and "." in texto:
        if texto.rfind(",") > texto.rfind("."):
            texto = texto.replace(".", "").replace(",", ".")
        else:
            texto = texto.replace(",", "")
    elif "," in texto:
        texto = texto.replace(",", ".")
    try:
        return float(texto)
    except ValueError:
        return None


def limpiar_order_id(valor):
    """ID_Pedido / tracking como texto limpio (evita notacion cientifica y '.0')."""
    if pd.isna(valor):
        return None
    texto = str(valor).strip()
    if texto.endswith(".0"):
        texto = texto[:-2]
    return texto if texto else None


def normalizar_pais(valor):
    """Aplica el mapa de paises. Si no esta en el mapa, deja el valor original."""
    if pd.isna(valor):
        return ""
    texto = str(valor).strip()
    return MAPA_PAIS.get(texto.lower(), texto)


def detectar_columna_tracking(columnas):
    """Busca la columna de seguimiento (Tracking ID, Tracking Number, ...)."""
    for c in columnas:
        if "tracking" in str(c).lower():
            return c
    return None


# ----------------------------------------------------------------------
# Lectura robusta de Excel
# ----------------------------------------------------------------------
def leer_excel_robusto(ruta):
    """
    TikTok manda el Excel de dos formas distintas:
      - Una "normal" que lee bien openpyxl.
      - Una "rara" en la que openpyxl solo ve la primera columna
        y hace falta el motor calamine.
    Esta funcion prueba ambos motores y se queda con el que
    devuelva mas de una columna (es decir, el que lee bien).
    """
    candidatos = []
    for motor in ("calamine", "openpyxl"):
        try:
            df = pd.read_excel(ruta, dtype=str, engine=motor)
            candidatos.append((df.shape[1], df))
        except Exception:
            pass

    if not candidatos:
        return pd.read_excel(ruta, dtype=str)

    candidatos.sort(key=lambda x: x[0], reverse=True)
    return candidatos[0][1]


# ----------------------------------------------------------------------
# Creacion de la carpeta de salida  (Cambio 1)
# ----------------------------------------------------------------------
def crear_carpeta_salida(base_dir, fecha_str):
    """Crea base_dir/YYYYMMDD_N usando el primer N libre (empezando en 1)."""
    n = 1
    while True:
        carpeta = base_dir / f"{fecha_str}_{n}"
        if not carpeta.exists():
            carpeta.mkdir(parents=True)
            return carpeta
        n += 1


# ----------------------------------------------------------------------
# Generacion de los manifiestos Word
# ----------------------------------------------------------------------
def _marcar_fila_cabecera(row):
    """Hace que esta fila de la tabla se repita como cabecera en cada pagina."""
    trPr = row._tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader")
    th.set(qn("w:val"), "true")
    trPr.append(th)


def _construir_cabecera(header, transportista, n_paquetes, fecha_corta):
    """Rellena la cabecera de pagina (se repite automaticamente en todas las hojas)."""
    # parrafo vacio inicial que trae el header por defecto
    header.paragraphs[0].text = ""

    # ---- Tabla 1 (sin bordes): transportista | MANIFIESTO + fecha ----
    t1 = header.add_table(rows=1, cols=2, width=Cm(17))
    izq, der = t1.cell(0, 0), t1.cell(0, 1)

    r_izq = izq.paragraphs[0].add_run(transportista)
    r_izq.bold = True
    r_izq.font.size = Pt(14)

    p_der1 = der.paragraphs[0]
    p_der1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_der1 = p_der1.add_run("MANIFIESTO DE ENTREGA")
    r_der1.bold = True
    r_der1.font.size = Pt(14)

    p_der2 = der.add_paragraph()
    p_der2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_der2.add_run(f"Fecha Entrega: {fecha_corta}")

    header.add_paragraph("")

    # ---- Tabla 2 (sin bordes): total | nombre transportista + firmas ----
    t2 = header.add_table(rows=1, cols=2, width=Cm(17))
    ci, cd = t2.cell(0, 0), t2.cell(0, 1)
    ci.paragraphs[0].add_run(f"Total de paquetes Entregados: {n_paquetes}")
    cd.paragraphs[0].add_run("Nombre del Transportista: ______________________")
    cd.add_paragraph("Firma Transportista:")
    cd.add_paragraph("Firma Vendedor:")

    # ---- Aviso (en rojo y negrita) ----
    aviso = header.add_paragraph()
    r_av = aviso.add_run("Por favor, firmar en todas las hojas")
    r_av.bold = True
    r_av.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)


def crear_manifiesto(ruta_doc, transportista, pares, fecha_corta):
    """
    Crea un manifiesto de entrega.

    transportista : "CTT" o "SEUR" (se muestra arriba a la izquierda).
    pares         : lista de tuplas (ID_Pedido, Tracking).
    fecha_corta   : texto DD/MM/YYYY.

    La cabecera (transportista, MANIFIESTO, fecha, total, nombre y firmas
    del transportista, firma del vendedor y el aviso) se coloca en la
    cabecera de pagina, de modo que se repite en TODAS las hojas.
    La fila de titulos de la tabla tambien se repite en cada pagina.
    """
    doc = Document()
    sec = doc.sections[0]
    # margen superior amplio para dejar sitio a la cabecera repetida
    sec.top_margin = Cm(5.8)
    sec.header_distance = Cm(0.8)

    _construir_cabecera(sec.header, transportista, len(pares), fecha_corta)

    # ---- Tabla de datos (CON bordes): Numero de Pedido + Numero de seguimiento ----
    tabla = doc.add_table(rows=1, cols=2)
    tabla.style = "Table Grid"
    hdr = tabla.rows[0].cells
    hdr[0].paragraphs[0].add_run("Numero de Pedido").bold = True
    hdr[1].paragraphs[0].add_run("Numero de seguimiento").bold = True
    _marcar_fila_cabecera(tabla.rows[0])   # repetir titulos en cada pagina

    for pid, track in pares:
        fila = tabla.add_row().cells
        fila[0].text = "" if pid is None else str(pid)
        fila[1].text = "" if track is None else str(track)

    doc.save(ruta_doc)


# ----------------------------------------------------------------------
# Programa principal
# ----------------------------------------------------------------------
def main():
    raiz = tk.Tk()
    raiz.withdraw()

    archivo = filedialog.askopenfilename(
        title="Selecciona el archivo descargado de TikTok",
        filetypes=[
            ("Excel o CSV", "*.xlsx *.xls *.csv"),
            ("Todos los archivos", "*.*"),
        ],
    )
    if not archivo:
        return

    if Document is None:
        messagebox.showerror(
            "Falta una libreria",
            "No se encuentra 'python-docx', necesaria para los manifiestos Word.\n\n"
            "Instalala con:  pip install python-docx",
        )
        return

    ruta = Path(archivo)

    # ---- Leer (todo como texto para no perder los ID de pedido) ----
    try:
        if ruta.suffix.lower() == ".csv":
            df = pd.read_csv(ruta, dtype=str, encoding="utf-8-sig")
        else:
            df = leer_excel_robusto(ruta)
    except Exception as e:
        messagebox.showerror("Error", f"No se pudo leer el archivo:\n{e}")
        return

    # ---- Quitar columnas vacias / "Unnamed" que arrastra Excel ----
    df = df.drop(columns=[c for c in df.columns if str(c).startswith("Unnamed")])
    df = df.dropna(axis=1, how="all")

    # ---- Quitar la segunda fila de cabecera (explicaciones) ----
    if len(df) > 0:
        valor = str(df.iloc[0].get("Order ID", "")).strip()
        if not re.fullmatch(r"\d{6,}", valor.replace(".0", "")):
            df = df.iloc[1:].reset_index(drop=True)

    # ---- Comprobar columnas ----
    faltan = [c for c in COLUMNAS_NECESARIAS if c not in df.columns]
    if faltan:
        # Caso tipico: se ha elegido un archivo ya procesado por este programa
        nombres_procesado = {"ID_Pedido", "Codigo_Producto", "Precio_Unitario"}
        if nombres_procesado.intersection(df.columns):
            messagebox.showerror(
                "Archivo equivocado",
                "Parece que has seleccionado un archivo YA PROCESADO "
                "(por ejemplo 'Orders_invoice.xlsx').\n\n"
                "Tienes que seleccionar el archivo ORIGINAL descargado de TikTok "
                "(el que tiene columnas en ingles: Order ID, Seller SKU, ...).",
            )
        else:
            messagebox.showerror(
                "Faltan columnas",
                "El archivo no tiene estas columnas:\n\n- " + "\n- ".join(faltan)
                + "\n\nColumnas encontradas:\n" + ", ".join(map(str, df.columns)),
            )
        return

    # ---- Detectar la columna de seguimiento (tracking) ----
    col_tracking = detectar_columna_tracking(df.columns)

    cols_a_usar = COLUMNAS_NECESARIAS + ([col_tracking] if col_tracking else [])
    df = df[cols_a_usar].copy()
    df = df.rename(columns=RENOMBRAR)
    if col_tracking:
        df = df.rename(columns={col_tracking: "Tracking"})
    else:
        df["Tracking"] = ""
    df["Tracking"] = df["Tracking"].apply(limpiar_order_id)

    filas_iniciales = len(df)

    # ---- Limpieza basica ----
    df["ID_Pedido"] = df["ID_Pedido"].apply(limpiar_order_id)
    df = df[df["ID_Pedido"].notna()].reset_index(drop=True)
    filas_sin_id = filas_iniciales - len(df)

    qty = pd.to_numeric(df["Qty"], errors="coerce")
    filas_sin_qty = int(qty.isna().sum())
    df["Qty"] = qty.astype("Int64")   # entero limpio para la factura

    df["Total"] = pd.to_numeric(df["Total"].map(limpiar_dinero), errors="coerce")

    # fecha real (para ordenar bien) y luego texto DD/MM/YYYY
    fecha_real = pd.to_datetime(df["Fecha_Recogida"], errors="coerce", dayfirst=False)
    df["Fecha_Recogida"] = fecha_real.dt.strftime("%d/%m/%Y")

    # ---- Normalizar pais (Cambio 3) ----
    df["Pais"] = df["Pais"].apply(normalizar_pais)

    # ---- Calculos: Precio_Unitario e IVA (Cambio 2) ----
    qty_safe = qty.replace(0, pd.NA)          # evita division por cero
    precio_unitario = df["Total"] / qty_safe / (1 + IVA_TIPO)
    df["Precio_Unitario"] = precio_unitario.round(2)
    df["IVA"] = (precio_unitario * IVA_TIPO).round(2)

    # ---- Ordenar por fecha, pais y pedido (Cambio 4) ----
    df["_orden_fecha"] = fecha_real
    df = df.sort_values(
        by=["_orden_fecha", "Pais", "ID_Pedido"], na_position="last"
    )
    df = df.drop(columns=["_orden_fecha"]).reset_index(drop=True)

    # ---- Factura final (Orders_invoice) ----
    factura = df[ORDEN_FINAL].copy()

    # ====================================================================
    # SALIDA
    # ====================================================================
    hoy = datetime.now()
    fecha_str = hoy.strftime("%Y%m%d")
    fecha_legible = f"{hoy.day} de {MESES_ES[hoy.month - 1]} de {hoy.year}"
    fecha_corta = hoy.strftime("%d/%m/%Y")   # para los manifiestos

    carpeta = crear_carpeta_salida(ruta.parent, fecha_str)

    orders_path = carpeta / "Orders_invoice.xlsx"
    resumen_path = carpeta / f"Resumen_{fecha_str}.xlsx"
    seur_path = carpeta / f"MANIFIESTO DE ENTREGA_SEUR_{fecha_str}.docx"
    ctt_path = carpeta / f"MANIFIESTO DE ENTREGA_CTT_{fecha_str}.docx"

    # ---- Orders_invoice.xlsx ----
    try:
        with pd.ExcelWriter(orders_path, engine="openpyxl") as writer:
            factura.to_excel(writer, index=False, sheet_name="Pedidos")
            hoja = writer.sheets["Pedidos"]
            # ID_Pedido es la columna B -> forzar texto
            for celda in hoja["B"]:
                celda.number_format = "@"
            # Qty (E) como entero
            for celda in hoja["E"][1:]:  # saltar cabecera
                celda.number_format = "0"
            # importes con 2 decimales (F=Precio_Unitario, G=IVA, H=Total)
            for col_letra in ("F", "G", "H"):
                for celda in hoja[col_letra][1:]:  # saltar cabecera
                    celda.number_format = "0.00"
    except PermissionError:
        messagebox.showerror(
            "Error",
            f"No se pudo guardar:\n{orders_path}\n\n"
            "Probablemente el archivo esta abierto en Excel. Cierralo y reintenta.",
        )
        return
    except Exception as e:
        messagebox.showerror("Error", f"No se pudo guardar Orders_invoice:\n{e}")
        return

    # ---- Calculo del resumen (Cambio 5) ----
    # Order Amount es a nivel de pedido: se toma una sola vez por ID_Pedido
    es_espana = df["Pais"] == "España"
    pedidos_esp = df[es_espana].drop_duplicates(subset="ID_Pedido")
    pedidos_ext = df[~es_espana].drop_duplicates(subset="ID_Pedido")

    esp_total = float(pedidos_esp["Total"].sum())
    ext_total = float(pedidos_ext["Total"].sum())
    esp_sin_iva = esp_total / (1 + IVA_TIPO)
    ext_sin_iva = ext_total / (1 + IVA_TIPO)
    esp_n = int(pedidos_esp["ID_Pedido"].nunique())
    ext_n = int(pedidos_ext["ID_Pedido"].nunique())

    filas_resumen = [
        ("España sin IVA", round(esp_sin_iva, 2)),
        ("España Total", round(esp_total, 2)),
        ("Numero de pedidos", esp_n),
        ("Extranjero sin IVA", round(ext_sin_iva, 2)),
        ("Extranjero Total", round(ext_total, 2)),
        ("Numero de pedidos", ext_n),
    ]

    try:
        from openpyxl import Workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "Resumen"
        for i, (etiqueta, valor) in enumerate(filas_resumen, start=1):
            ws.cell(row=i, column=1, value=etiqueta)
            ws.cell(row=i, column=2, value=valor)
        wb.save(resumen_path)
    except Exception as e:
        messagebox.showerror("Error", f"No se pudo guardar el Resumen:\n{e}")
        return

    # ---- Manifiestos Word ----
    # df ya esta ordenado; drop_duplicates conserva ese orden.
    # Un pedido = un paquete -> una fila por ID_Pedido con su tracking.
    man_esp = (df.loc[es_espana, ["ID_Pedido", "Tracking"]]
                 .drop_duplicates(subset="ID_Pedido"))
    man_ext = (df.loc[~es_espana, ["ID_Pedido", "Tracking"]]
                 .drop_duplicates(subset="ID_Pedido"))
    pares_espana = list(man_esp.itertuples(index=False, name=None))
    pares_extranjero = list(man_ext.itertuples(index=False, name=None))

    try:
        # SEUR -> extranjero (NO Espana)
        crear_manifiesto(seur_path, "SEUR", pares_extranjero, fecha_corta)
        # CTT -> Espana
        crear_manifiesto(ctt_path, "CTT", pares_espana, fecha_corta)
    except Exception as e:
        messagebox.showerror("Error", f"No se pudieron crear los manifiestos:\n{e}")
        return

    # ---- Resumen final ----
    mensaje = (
        f"Carpeta creada:\n{carpeta}\n\n"
        f"Archivos generados:\n"
        f"  - {orders_path.name}\n"
        f"  - {resumen_path.name}\n"
        f"  - {seur_path.name}\n"
        f"  - {ctt_path.name}\n\n"
        f"Lineas procesadas: {len(factura)}\n"
        f"Pedidos España: {esp_n}    Pedidos extranjero: {ext_n}\n"
        f"Filas eliminadas sin ID_Pedido: {filas_sin_id}"
    )
    if filas_sin_qty > 0:
        mensaje += (
            f"\n\nATENCION: {filas_sin_qty} fila(s) sin Qty. "
            "Revisalas (el Precio_Unitario no se pudo calcular)."
        )
    if not col_tracking:
        mensaje += (
            "\n\nATENCION: no se encontro la columna de seguimiento (Tracking) "
            "en el archivo; los manifiestos saldran SIN numero de seguimiento."
        )
    messagebox.showinfo("Listo", mensaje)


if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        root = tk.Tk()
        root.withdraw()
        messagebox.showerror("Error inesperado", str(e))
        sys.exit(1)